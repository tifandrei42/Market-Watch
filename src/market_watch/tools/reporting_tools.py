import yfinance as yf
import pandas as pd
import pandas_ta as ta
import matplotlib.pyplot as plt
import os
import json
import re
from crewai.tools import BaseTool
from typing import Type, List
from pydantic import BaseModel, Field
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class WordReportToolInput(BaseModel):
    report_content: str = Field(..., description="The full markdown content of the report.")
    chart_paths: List[str] = Field(default=[], description="List of file paths to the generated chart PNGs.")

class WordReportTool(BaseTool):
    name: str = "Word Report Tool"
    description: str = (
        "Generates a professional Word Document (.docx) from the provided markdown content and embeds "
        "any specified chart images. Saves the file to 'output/market_watch_report.docx'. "
        "Also generates a 'dashboard_data.json' file for the frontend."
    )
    args_schema: Type[BaseModel] = WordReportToolInput

    def _run(self, report_content: str, chart_paths: List[str] = []) -> str:
        try:
            # --- DOCX GENERATION ---
            doc = Document()
            title = doc.add_heading('Market Watch Daily Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for line in report_content.split('\n'):
                if line.startswith('# '):
                   doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                   doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                   doc.add_heading(line[4:], level=3)
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:-2])
                    run.bold = True
                elif line.strip() == '---':
                    doc.add_page_break()
                else:
                    doc.add_paragraph(line)

            if chart_paths:
                doc.add_heading('Market Visuals', level=1)
                for path in chart_paths:
                    if os.path.exists(path):
                        doc.add_picture(path, width=Inches(6))
                        last_paragraph = doc.paragraphs[-1] 
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph(f"Figure: {os.path.basename(path)}", style='Caption')
                    else:
                         doc.add_paragraph(f"[Missing Image: {path}]", style='Quote')

            output_dir = "output"
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            
            docx_path = os.path.join(output_dir, "market_watch_report.docx")
            doc.save(docx_path)

            # --- JSON DATA EXTRACTION ---
            top_short = []
            top_long = []
            current_section = None

            for line in report_content.split('\n'):
                line = line.strip()
                if "Top 5 Short-Term Picks" in line:
                    current_section = "short"
                    continue
                elif "Top 5 Long-Term Picks" in line:
                    current_section = "long"
                    continue
                elif line.startswith("#"):
                    current_section = None
                
                if current_section and line.startswith("-"):
                    # Extract Ticker and Reason
                    # Expected format: "- **NVDA**: Reason..." or "- NVDA: Reason..."
                    match = re.search(r"\- (?:\*\*)?([A-Z]+)(?:\*\*)?[:\s]+(.*)", line)
                    if match:
                        pick = {"ticker": match.group(1), "reason": match.group(2)}
                        if current_section == "short":
                            top_short.append(pick)
                        elif current_section == "long":
                            top_long.append(pick)

            dashboard_data = {
                "generated_at": pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S'),
                "top_short": top_short,
                "top_long": top_long,
                "charts": [os.path.basename(p) for p in chart_paths],
                "full_report": report_content
            }

            json_path = os.path.join(output_dir, "dashboard_data.json")
            with open(json_path, 'w') as f:
                json.dump(dashboard_data, f, indent=2)
            
            return f"Report saved to {docx_path} and data to {json_path}"

        except Exception as e:
            return f"Error creating reports: {str(e)}"
